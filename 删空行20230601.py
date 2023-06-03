import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import Cm

def process_word_files(folder_path):

        for file in os.listdir(folder_path):
            if file.endswith(".docx"):
                file_path = os.path.join(folder_path, file)
                update_word_file(file_path)

def update_word_file(file_path):
        document = Document(file_path)

        top_margin = Cm(3.7)
        bottom_margin = Cm(3.5)
        left_margin = Cm(2.8)
        right_margin = Cm(2.6)

        paragraphs_to_remove = []

        sections = document.sections #设置页边距
        for section in sections:
            section.top_margin = top_margin
            section.bottom_margin = bottom_margin
            section.left_margin = left_margin
            section.right_margin = right_margin

        for index, paragraph in enumerate(document.paragraphs):
            if len(paragraph.runs) == 0:
                paragraphs_to_remove.append(paragraph)
                continue

            if index == 0:

                paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER #第一段居中

                run = paragraph.runs[0]
                run.font.name = '方正小标宋简体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')
                run.font.size = Pt(22)
            elif index in [1, 2]:

                paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER #第二三段居中

                for run in paragraph.runs:
                    run.font.name = '仿宋_GB2312'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
                    run.font.size = Pt(16)
                    run.font.bold = True


            else:
                for run in paragraph.runs:
                    run.font.name = '仿宋_GB2312'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
                    run.font.size = Pt(16)
                    run.text = run.text.strip() #去掉空格
                    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT #第四段开始以后居左
                    paragraph.paragraph_format.line_spacing = Pt(29.8) #行间距
                    paragraph.paragraph_format.first_line_indent = Pt(32) #首行缩进

                if paragraph.text.isspace() or paragraph.text == "": #去掉空行
                    paragraphs_to_remove.append(paragraph)

        for paragraph in paragraphs_to_remove:
            p = paragraph._element
            p.getparent().remove(p)

        document.paragraphs[3].insert_paragraph_before('') #第三段后空两行
        document.paragraphs[3].insert_paragraph_before('')

        document.save(file_path)



if __name__ == "__main__":
    folder_path = "D:/1/test"  # 请将此路径替换为实际的文件夹路径
    process_word_files(folder_path)
